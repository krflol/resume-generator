from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def add_section(document, text, style):
    section = document.add_paragraph(text, style=style)
    section.paragraph_format.space_after = Pt(12)

def add_bullet_points(document, bullet_points, style):
    for point in bullet_points:
        document.add_paragraph(point, style=style)
        document.paragraphs[-1].paragraph_format.left_indent = Pt(36)

# Create a new Word document
document = Document()

# Define styles
document.styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
document.styles['Name'].font.size = Pt(18)
document.styles['Name'].font.bold = True

document.styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
document.styles['Contact'].font.size = Pt(11)

document.styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
document.styles['Section'].font.size = Pt(14)
document.styles['Section'].font.bold = True

document.styles.add_style('Skill', WD_STYLE_TYPE.PARAGRAPH)
document.styles['Skill'].font.size = Pt(11)

document.styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
document.styles['JobTitle'].font.size = Pt(12)
document.styles['JobTitle'].font.bold = True

document.styles.add_style('BulletPoint', WD_STYLE_TYPE.PARAGRAPH)
document.styles['BulletPoint'].font.size = Pt(11)

# Add name and contact information
add_section(document, 'Keenan Finkelstein', 'Name')
add_section(document, 'Pensacola, FL', 'Contact')
add_section(document, 'keenanfinkelstein@gmail.com', 'Contact')
add_section(document, '850-619-0756', 'Contact')
add_section(document, 'linkedin.com/in/keenan-finkelstein-96200a42/', 'Contact')
add_section(document, 'https://github.com/krflol/brAInstorm', 'Contact')
add_section(document, 'https://github.com/krflol/Adornment', 'Contact')

# Add summary
add_section(document, 'Software Engineer', 'Section')
add_section(document, 'Accomplished Software Engineer with over a decade of experience delivering innovative solutions for global companies across diverse industries, including publishing, analytics, laboratories, direct mailing, and the U.S. Government. Proven expertise in full-stack development, digital workflows, automation, and RESTful APIs. Consistently improves systems to ensure operational excellence, drives significant revenue growth through strategic feature development, and enables comprehensive data collection to ensure compliance with federal regulations.', 'Skill')

# Add technical skills
add_section(document, 'Technical Skills', 'Section')
add_bullet_points(document, [
    'Languages: Python, Perl, HTML, CSS, SQL, VBA, C# (WPF)',
    'Frameworks/Libraries: React, Flask, Django, FastAPI, Pandas, NumPy, TensorFlow, OpenAI, PyTorch, Selenium, BeautifulSoup, Requests',
    'Tools: AWS, Adobe Creative Suite (InDesign, Illustrator, Photoshop)',
    'Methodologies: Agile, Test Automation, Machine Learning, Deep Reinforcement Learning',
    'Other: RESTful APIs, Solution Design, Project Support'
], 'Skill')

# Add professional experience
add_section(document, 'Professional Experience', 'Section')

add_section(document, 'Software Engineer IV - Elsevier (Remote, 2022-2023)', 'JobTitle')
add_bullet_points(document, [
    'Skillfully maintained and optimized a complex 20-year-old Perl monolith serving as the critical document repository for academic journals',
    'Spearheaded the development and launch of a Django application for efficient bulk uploads',
    'Developed and implemented high-impact helper scripts for the team using Perl and Python, streamlining workflows and boosting productivity'
], 'BulletPoint')

add_section(document, 'Senior Software Engineer and IT Manager - Regenative Labs (Pensacola, FL, 2020-2022)', 'JobTitle')
add_bullet_points(document, [
    'Engineered and maintained critical features for patient data collection, sales, purchasing, and distribution of cutting-edge medical products, directly contributing to explosive monthly revenue growth from hundreds of thousands to millions of dollars',
    'Pioneered a "page as component" design pattern using React to rapidly implement business-critical features while minimizing interaction with legacy codebase, resulting in significant efficiency gains',
    'Architected and built a RESTful API as a backend to the React frontend, enabling near real-time feature implementation and enhancing system responsiveness',
    'Designed and deployed intuitive dashboards empowering healthcare providers and staff to effectively use and observe submitted information, greatly improving data visibility and decision-making',
    'Developed a game-changing digital "Prior Authorization" workflow with real-time communication between providers and billers, streamlining the authorization process and reducing delays',
    'Successfully brought the project into full HIPAA compliance, enabling secure data collection while maintaining strict regulatory standards',
    'Mentored and inspired junior engineers in React, HTML, CSS, Rust, and Python, fostering a culture of continuous learning and growth'
], 'BulletPoint')

add_section(document, 'Senior Software Engineer and Manager - Action Mailing Service (Pensacola, FL, 2017-2020)', 'JobTitle')
add_bullet_points(document, [
    'Spearheaded the development of custom solutions for a high-volume mailing company using a VBA tech stack, with a strong emphasis on automation and efficiency',
    'Implemented a comprehensive modernization of the technology stack through innovative automated solutions in Python and C# (WPF), resulting in significant performance improvements',
    'Forged strong collaborative relationships with clients and colleagues to ensure optimal project delivery and client satisfaction'
], 'BulletPoint')

# Add education and certifications
add_section(document, 'Education and Certifications', 'Section')
add_section(document, 'Earned certifications and completed extensive training in submarine electronics and sonar technologies from the U.S. Navy', 'Skill')

# Add additional experience
add_section(document, 'Additional Experience', 'Section')
add_bullet_points(document, [
    'Computer Field/Sonar Technician, U.S. Navy',
    'Installation Technician, DirecTV',
    'Vacation Rental Maintenance Technician, Property Care Solutions'
], 'BulletPoint')

# Save the document
document.save('KeenanFinkelstein_Resume.docx')